# core/services/document_processor.py
import PyPDF2
import docx
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
import re
from pathlib import Path
import cv2
import numpy as np



def extract_text_with_ocr(pdf_path):
    """Extract text using OCR for scanned PDFs."""
    images = convert_from_path(pdf_path)
    extracted_text = ""
    
    for img in images:
        img_cv = np.array(img)
        img_gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)  # Convert to grayscale
        extracted_text += pytesseract.image_to_string(img_gray, config='--psm 6') + "\n"

    return extracted_text.strip()

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF files using PyPDF2 and fallback to OCR if needed"""
    text = ""
    try:
        # Try direct text extraction first
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text() + "\n"
                
        # If extracted text is too little, use OCR
        # if len(text.strip()) < 100:  # Arbitrary threshold
        #     images = pdf2image.convert_from_path(pdf_path)
        #     for img in images:
        #         text += pytesseract.image_to_string(img) + "\n"
        # If no text found, try OCR
        if not text.strip():
            text = extract_text_with_ocr(pdf_path)

    except Exception as e:
        print(f"Error processing PDF {pdf_path}: {e}")
    return text



def extract_text_from_docx(docx_path):
    """Extract text from Word documents, including tables, paragraphs, and text boxes."""
    doc = docx.Document(docx_path)
    text = ""

    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Avoid empty lines
            text += para.text + "\n"

    # Extract text from tables (for column-like content)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():  # Avoid empty cells
                    text += cell.text + "\n"

    # Extract text from text boxes (for column-like content)
    for shape in doc.inline_shapes:
        if shape.type == 3:  # Check if it's a text box (type 3 is text box)
            text += shape.text_frame.text + "\n"

    return text.strip()

def preprocess_text(text):
    """Clean and preprocess extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    # Remove special characters but keep relevant punctuation
    text = re.sub(r'[^\w\s.,@:;()/-]', '', text)
    return text

def process_document(file_path):
    """Process a document based on its file extension"""
    file_path = Path(file_path)
    
    if file_path.suffix.lower() == '.pdf':
        raw_text = extract_text_from_pdf(file_path)
    elif file_path.suffix.lower() in ['.docx', '.doc']:
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_path.suffix}")
    
    return preprocess_text(raw_text)